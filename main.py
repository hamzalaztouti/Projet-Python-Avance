"""
main.py - Partie 2 : Analyse textuelle et génération du rapport Word
Projet Python Avancé - Hamza Laztouti / Tahiri Rossaina
Source : Les Misérables, Victor Hugo - Project Gutenberg
"""
# IMPORTS
import urllib.request
import matplotlib.pyplot as plt
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# COULEURS
BLEU_FONCE  = "1A2D5A"
ROUGE_VIF   = "C0392B"
BLEU_CLAIR  = "2E5FA3"
GRIS_TEXTE  = "3D3D3D"
BLANC       = "FFFFFF"

# ÉTAPE 1 : TÉLÉCHARGEMENT DU LIVRE

def telecharger_livre(url, chemin_local):
    """Télécharge le fichier texte du livre depuis Gutenberg."""
    print("Téléchargement en cours...")
    try:
        urllib.request.urlretrieve(url, chemin_local)
        print("Téléchargement terminé !")
    except Exception as e:
        print("Erreur lors du téléchargement du livre : " + str(e))
        exit()


def lire_livre(chemin_local):
    """Lit le fichier texte et retourne la liste de lignes."""
    try:
        fichier = open(chemin_local, "r", encoding="utf-8")
        contenu = fichier.read()
        fichier.close()
        return contenu.split("\n")
    except Exception as e:
        print("Erreur lors de la lecture du fichier : " + str(e))
        exit()


def extraire_titre_auteur(lignes):
    """Extrait le titre et l'auteur depuis les métadonnées Gutenberg."""
    titre = ""
    auteur = ""
    for ligne in lignes:
        if ligne.startswith("Title:"):
            titre = ligne.replace("Title:", "").strip()
        if ligne.startswith("Author:"):
            auteur = ligne.replace("Author:", "").strip()
        if titre != "" and auteur != "":
            break
    return titre, auteur


def extraire_premier_chapitre(lignes):
    """Extrait le texte du premier chapitre (entre Chapitre I et Chapitre II)."""
    premier_chapitre = ""
    dans_chapitre = False
    for i in range(len(lignes)):
        ligne = lignes[i].strip()
        if ligne == "Chapitre I" and i > 60:
            dans_chapitre = True
            continue
        if dans_chapitre:
            if ligne == "Chapitre II":
                break
            premier_chapitre = premier_chapitre + lignes[i] + "\n"
    return premier_chapitre

# ÉTAPE 2 : ANALYSE DES PARAGRAPHES

def compter_mots_paragraphes(texte_chapitre):
    """Compte les mots par paragraphe, arrondit à la dizaine, retourne la liste triée."""
    try:
        paragraphes = texte_chapitre.split("\n\n")
        nb_mots_par_paragraphe = []
        for para in paragraphes:
            para = para.strip()
            if para != "":
                mots = para.split()
                nb_mots = len(mots)
                nb_mots_arrondi = (nb_mots // 10) * 10
                if nb_mots_arrondi > 0:
                    nb_mots_par_paragraphe.append(nb_mots_arrondi)
        nb_mots_par_paragraphe.sort()
        return nb_mots_par_paragraphe
    except Exception as e:
        print("Erreur lors du comptage des mots : " + str(e))
        exit()


def calculer_occurrences(nb_mots_par_paragraphe):
    """Calcule les longueurs uniques et leur nombre d'occurrences."""
    longueurs = []
    occurrences = []
    for valeur in nb_mots_par_paragraphe:
        if valeur not in longueurs:
            longueurs.append(valeur)
            occurrences.append(nb_mots_par_paragraphe.count(valeur))
    return longueurs, occurrences


# ÉTAPE 3 : GRAPHIQUE

def generer_graphique(longueurs, occurrences, chemin_sortie):
    """Génère et sauvegarde le graphique de distribution."""
    try:
        fig, ax = plt.subplots(figsize=(12, 6))
        ax.bar(longueurs, occurrences, width=8, color="#1A2D5A", edgecolor="#C0392B", linewidth=0.8)
        ax.set_title("Distribution des longueurs des paragraphes — Chapitre I",
                     fontsize=14, fontweight="bold", color="#1A2D5A", pad=15)
        ax.set_xlabel("Nombre de mots (arrondi à la dizaine)", fontsize=11, color="#3D3D3D")
        ax.set_ylabel("Nombre de paragraphes", fontsize=11, color="#3D3D3D")
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.spines["left"].set_color("#CCCCCC")
        ax.spines["bottom"].set_color("#CCCCCC")
        ax.tick_params(colors="#3D3D3D")
        ax.grid(axis="y", color="#EEEEEE", linewidth=0.8)
        plt.tight_layout()
        plt.savefig(chemin_sortie, dpi=150)
        plt.close()
        print("Graphique sauvegardé : " + chemin_sortie)
    except Exception as e:
        print("Erreur lors de la génération du graphique : " + str(e))

# ÉTAPE 4 : IMAGE

def telecharger_image(url_image, chemin_sortie):
    """Télécharge l'image de couverture depuis Internet."""
    print("Téléchargement image en cours...")
    try:
        requete = urllib.request.Request(url_image)
        requete.add_header("User-Agent", "Mozilla/5.0")
        reponse = urllib.request.urlopen(requete)
        fichier_image = open(chemin_sortie, "wb")
        fichier_image.write(reponse.read())
        fichier_image.close()
        print("Image téléchargée : " + chemin_sortie)
    except Exception as e:
        print("Erreur téléchargement image : " + str(e))


def traiter_image(chemin_image, chemin_sortie):
    """Recadre et redimensionne l'image originale."""
    try:
        img1 = Image.open(chemin_image)
        print("Taille originale : " + str(img1.size))
        largeur, hauteur = img1.size
        gauche = largeur // 4
        haut = hauteur // 4
        droite = largeur * 3 // 4
        bas = hauteur * 3 // 4
        img1_recadree = img1
        img1_finale = img1_recadree.resize((400, 500))
        img1_finale.save(chemin_sortie)
        print("Image recadrée et redimensionnée : " + str(img1_finale.size))
    except Exception as e:
        print("Erreur lors du traitement de l'image : " + str(e))


def creer_logo(chemin_sortie):
    """Crée un logo simple noir et blanc avec les initiales VH."""
    try:
        logo = Image.new("L", (100, 100), color=255)
        dessin = ImageDraw.Draw(logo)
        dessin.ellipse((10, 10, 90, 90), outline=0, width=5)
        dessin.text((30, 35), "VH", fill=0)
        logo.save(chemin_sortie)
        print("Logo créé : " + chemin_sortie)
        return logo
    except Exception as e:
        print("Erreur lors de la création du logo : " + str(e))


def coller_logo_sur_image(chemin_image, logo, chemin_sortie, angle=45):
    """Fait pivoter le logo et le colle sur l'image principale."""
    try:
        logo_pivote = logo.rotate(angle)
        img = Image.open(chemin_image).convert("RGB")
        logo_rgb = logo_pivote.convert("RGB")
        img.paste(logo_rgb, (10, 10))
        img.save(chemin_sortie)
        print("Logo collé sur image : " + chemin_sortie)
    except Exception as e:
        print("Erreur lors du collage du logo : " + str(e))

# ÉTAPE 5 : RAPPORT WORD

def set_fond_paragraphe(paragraph, couleur_hex):
    """Applique une couleur de fond à un paragraphe."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), couleur_hex)
    pPr.append(shd)


def ajouter_ligne_coloree(doc, couleur_hex, epaisseur=24):
    """Ajoute une ligne horizontale colorée sous un paragraphe."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(epaisseur))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), couleur_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def para_centre(doc, texte, taille, gras=False, italique=False,
                couleur=None, police="Times New Roman",
                espace_avant=0, espace_apres=0, fond=None):
    """Crée un paragraphe centré avec style."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(espace_avant)
    p.paragraph_format.space_after = Pt(espace_apres)
    if fond:
        set_fond_paragraphe(p, fond)
    run = p.add_run(texte)
    run.bold = gras
    run.italic = italique
    run.font.size = Pt(taille)
    run.font.name = police
    if couleur:
        r, g, b = int(couleur[0:2], 16), int(couleur[2:4], 16), int(couleur[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    return p


def generer_rapport_word(titre, auteur, nb_mots_par_paragraphe, chemin_sortie):
    """Génère le rapport Word complet professionnel bleu/rouge Times New Roman."""
    try:
        doc = Document()
        nb_total_mots = sum(nb_mots_par_paragraphe)

        # Marges page de titre
        section = doc.sections[0]
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)

        # BANDE BLEUE HAUT
        for _ in range(4):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_fond_paragraphe(p, BLEU_FONCE)

        # Titre principal
        para_centre(doc, titre, 30, gras=True, couleur=BLANC,
                    police="Times New Roman", espace_avant=20, espace_apres=8,
                    fond=BLEU_FONCE)

        # Sous-titre
        para_centre(doc, "Analyse littéraire  —  Projet Python Avancé", 13,
                    italique=True, couleur="A8C0E0",
                    police="Times New Roman", espace_avant=0, espace_apres=20,
                    fond=BLEU_FONCE)

        # Bas bande bleue
        for _ in range(3):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_fond_paragraphe(p, BLEU_FONCE)

        # Ligne rouge décorative
        ajouter_ligne_coloree(doc, ROUGE_VIF, 32)

        # Espace
        doc.add_paragraph().paragraph_format.space_after = Pt(16)

        # Image centrée
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after = Pt(16)
        p_img.paragraph_format.left_indent = Cm(2.5)
        p_img.paragraph_format.right_indent = Cm(2.5)
        run_img = p_img.add_run()
        run_img.add_picture("image1_avec_logo.jpg", width=Inches(2.6))

        # Ligne rouge sous image
        p_ligne = doc.add_paragraph()
        p_ligne.paragraph_format.left_indent = Cm(2.5)
        p_ligne.paragraph_format.right_indent = Cm(2.5)
        p_ligne.paragraph_format.space_before = Pt(0)
        p_ligne.paragraph_format.space_after = Pt(16)
        pPr = p_ligne._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "16")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), ROUGE_VIF)
        pBdr.append(bottom)
        pPr.append(pBdr)

        # Auteur du livre
        p_auteur = doc.add_paragraph()
        p_auteur.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_auteur.paragraph_format.left_indent = Cm(2.5)
        p_auteur.paragraph_format.right_indent = Cm(2.5)
        p_auteur.paragraph_format.space_before = Pt(8)
        p_auteur.paragraph_format.space_after = Pt(6)
        run_lbl1 = p_auteur.add_run("Auteur de l'œuvre  ")
        run_lbl1.font.size = Pt(11)
        run_lbl1.font.name = "Times New Roman"
        run_lbl1.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run_val1 = p_auteur.add_run(auteur)
        run_val1.bold = True
        run_val1.font.size = Pt(15)
        run_val1.font.name = "Times New Roman"
        run_val1.font.color.rgb = RGBColor(0x1A, 0x2D, 0x5A)

        # Auteurs du rapport
        p_rapport = doc.add_paragraph()
        p_rapport.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rapport.paragraph_format.left_indent = Cm(2.5)
        p_rapport.paragraph_format.right_indent = Cm(2.5)
        p_rapport.paragraph_format.space_before = Pt(6)
        p_rapport.paragraph_format.space_after = Pt(16)
        run_lbl2 = p_rapport.add_run("Rapport rédigé par  ")
        run_lbl2.font.size = Pt(11)
        run_lbl2.font.name = "Times New Roman"
        run_lbl2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run_val2 = p_rapport.add_run("Hamza Laztouti  &  Tahiri Rossaina")
        run_val2.bold = True
        run_val2.font.size = Pt(15)
        run_val2.font.name = "Times New Roman"
        run_val2.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

        # Bande bleue bas
        ajouter_ligne_coloree(doc, BLEU_FONCE, 24)

        # Source
        p_src = doc.add_paragraph()
        p_src.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_src.paragraph_format.left_indent = Cm(2.5)
        p_src.paragraph_format.right_indent = Cm(2.5)
        p_src.paragraph_format.space_before = Pt(10)
        run_src = p_src.add_run("Source : Project Gutenberg  •  gutenberg.org")
        run_src.italic = True
        run_src.font.size = Pt(9)
        run_src.font.name = "Times New Roman"
        run_src.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        doc.add_page_break()

        # Marges normales pages suivantes
        section2 = doc.add_section()
        section2.top_margin = Cm(2.5)
        section2.bottom_margin = Cm(2.5)
        section2.left_margin = Cm(2.5)
        section2.right_margin = Cm(2.5)

        # PAGE GRAPHIQUE
        h1 = doc.add_heading("Analyse du Chapitre I", level=1)
        h1.runs[0].font.color.rgb = RGBColor(0x1A, 0x2D, 0x5A)
        h1.runs[0].font.name = "Times New Roman"

        h2 = doc.add_heading("Distribution des longueurs de paragraphes", level=2)
        h2.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        h2.runs[0].font.name = "Times New Roman"

        doc.add_picture("graphique_chapitre1.png", width=Inches(5.5))

        p_desc_label = doc.add_paragraph()
        run_dl = p_desc_label.add_run("Description du graphique")
        run_dl.bold = True
        run_dl.italic = True
        run_dl.font.size = Pt(12)
        run_dl.font.name = "Times New Roman"
        run_dl.font.color.rgb = RGBColor(0x1A, 0x2D, 0x5A)

        description = ("Ce graphique montre la distribution des longueurs des paragraphes du Chapitre I. "
                       "Nombre de paragraphes : " + str(len(nb_mots_par_paragraphe)) + ". "
                       "Nombre total de mots : " + str(nb_total_mots) + ". "
                       "Minimum : " + str(min(nb_mots_par_paragraphe)) + " mots. "
                       "Maximum : " + str(max(nb_mots_par_paragraphe)) + " mots. "
                       "Moyenne : " + str(nb_total_mots // len(nb_mots_par_paragraphe)) + " mots. "
                       "Source : Project Gutenberg.")
        p_desc = doc.add_paragraph(description)
        p_desc.runs[0].font.size = Pt(11)
        p_desc.runs[0].font.name = "Times New Roman"

        doc.add_page_break()

        # PAGE STATISTIQUES
        h3 = doc.add_heading("Statistiques détaillées", level=1)
        h3.runs[0].font.color.rgb = RGBColor(0x1A, 0x2D, 0x5A)
        h3.runs[0].font.name = "Times New Roman"

        h4 = doc.add_heading("Source des données", level=2)
        h4.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        h4.runs[0].font.name = "Times New Roman"

        p_source = doc.add_paragraph()
        r1 = p_source.add_run("Livre téléchargé depuis ")
        r1.font.size = Pt(11)
        r1.font.name = "Times New Roman"
        r2 = p_source.add_run("Project Gutenberg (gutenberg.org)")
        r2.bold = True
        r2.font.size = Pt(11)
        r2.font.name = "Times New Roman"

        h5 = doc.add_heading("Chiffres clés du Chapitre I", level=2)
        h5.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        h5.runs[0].font.name = "Times New Roman"

        chiffres = [
            ("Nombre de paragraphes analysés", str(len(nb_mots_par_paragraphe))),
            ("Nombre total de mots",           str(nb_total_mots)),
            ("Paragraphe le plus court",       str(min(nb_mots_par_paragraphe)) + " mots"),
            ("Paragraphe le plus long",        str(max(nb_mots_par_paragraphe)) + " mots"),
            ("Moyenne de mots par paragraphe", str(nb_total_mots // len(nb_mots_par_paragraphe)) + " mots"),
        ]
        for label, valeur in chiffres:
            p = doc.add_paragraph()
            rl = p.add_run(label + " : ")
            rl.bold = True
            rl.font.size = Pt(11)
            rl.font.name = "Times New Roman"
            rl.font.color.rgb = RGBColor(0x1A, 0x2D, 0x5A)
            rv = p.add_run(valeur)
            rv.italic = True
            rv.font.size = Pt(11)
            rv.font.name = "Times New Roman"
            rv.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

        h6 = doc.add_heading("Note méthodologique", level=2)
        h6.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        h6.runs[0].font.name = "Times New Roman"

        p_note = doc.add_paragraph()
        rn = p_note.add_run(
            "Les nombres de mots sont arrondis à la dizaine inférieure "
            "(ex : 127 → 120). Les paragraphes de moins de 10 mots et les "
            "paragraphes vides sont exclus de l'analyse."
        )
        rn.italic = True
        rn.font.size = Pt(10)
        rn.font.name = "Times New Roman"
        rn.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        doc.save(chemin_sortie)
        print("Document Word sauvegardé : " + chemin_sortie)

    except Exception as e:
        print("Erreur lors de la génération du rapport Word : " + str(e))


# PROGRAMME PRINCIPAL
if __name__ == "__main__":

    URL_LIVRE         = "https://www.gutenberg.org/cache/epub/17489/pg17489.txt"
    URL_IMAGE = "https://www.gutenberg.org/cache/epub/135/pg135.cover.medium.jpg"

    CHEMIN_LIVRE      = "les_miserables.txt"
    CHEMIN_IMAGE      = "image12.jpg"
    CHEMIN_IMAGE_MOD  = "image1_modifiee.jpg"
    CHEMIN_IMAGE_LOGO = "image1_avec_logo.jpg"
    CHEMIN_LOGO       = "logo.png"
    CHEMIN_GRAPHIQUE  = "graphique_chapitre1.png"
    CHEMIN_RAPPORT    = "rapport_les_miserables.docx"

    # Étape 1 : Téléchargement et lecture du livre
    telecharger_livre(URL_LIVRE, CHEMIN_LIVRE)
    lignes = lire_livre(CHEMIN_LIVRE)
    titre, auteur = extraire_titre_auteur(lignes)
    print("Titre : " + titre)
    print("Auteur : " + auteur)

    # Étape 2 : Extraction et analyse du chapitre
    premier_chapitre = extraire_premier_chapitre(lignes)
    print("Premier chapitre extrait : " + str(len(premier_chapitre)) + " caractères")

    nb_mots_par_paragraphe = compter_mots_paragraphes(premier_chapitre)
    print("Nombre de paragraphes : " + str(len(nb_mots_par_paragraphe)))

    longueurs, occurrences = calculer_occurrences(nb_mots_par_paragraphe)

    # Étape 3 : Graphique
    generer_graphique(longueurs, occurrences, CHEMIN_GRAPHIQUE)

    # Étape 4 : Image

    # telecharger_image(URL_IMAGE, CHEMIN_IMAGE)
    traiter_image(CHEMIN_IMAGE, CHEMIN_IMAGE_MOD)
    logo = creer_logo(CHEMIN_LOGO)
    coller_logo_sur_image(CHEMIN_IMAGE_MOD, logo, CHEMIN_IMAGE_LOGO)

    # Étape 5 : Rapport Word
    generer_rapport_word(titre, auteur, nb_mots_par_paragraphe, CHEMIN_RAPPORT)